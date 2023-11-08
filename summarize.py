from langchain.chains.combine_documents.stuff import StuffDocumentsChain
from langchain.docstore.document import Document
from langchain.prompts import PromptTemplate
from langchain.chains.llm import LLMChain
from utils import convert_json, get_items
from langdetect import detect_langs
from tqdm import tqdm
import json

class generate_summary():
    def __init__(self, file_name:str, original_language:str, translated_language:str, chunk_size:int, output_dir:str) -> None:
        from utils import llm
        self.file_name = file_name
        self.chunk_size = chunk_size 
        self.original_language = original_language
        self.translated_language = translated_language
        self.output_dir = output_dir
        self.llm = llm

    def _get_general_summary(self, article_divided:dict) -> None:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        from langchain.chains.summarize import load_summarize_chain

        text_splitter = RecursiveCharacterTextSplitter(
                            chunk_size = self.chunk_size//2,
                            chunk_overlap  = 0,
                            length_function = len)
        
        # load transcript
        with open(f'./transcript/{self.file_name}.txt', 'r') as f:
            transcript = ''.join(f.readlines())
        split_text = text_splitter.split_text(transcript)

        item_list, items, item_format  = get_items('general')
        prompt_template = f"###Write cosine {items} of the following:###" "{text} \n"
        prompt = PromptTemplate.from_template(prompt_template)

        refine_template = (
            f"Your job is to produce a final streamline {items}.\
            We have provided an existing {items} up to a certain point:""{existing_answer}\n"\
            f"We have the opportunity to refine {items}"
            "(only if needed) with some more context below.\n\
            ------------\n\
            {text}\n\
            ------------\n"
            f"Given the new context, refine the original {items} in {self.original_language}\
            If the context isn't useful, return the origin {items} in {self.original_language}\
            Fulfill the format: {item_format}"
    )
        refine_prompt = PromptTemplate.from_template(refine_template)

        chain = load_summarize_chain(
            llm=self.llm,
            chain_type="refine",
            question_prompt=prompt,
            refine_prompt=refine_prompt,
            return_intermediate_steps=False,
            input_key="input_documents",
            output_key="output_text"
        )
        print('Analysing general items')
        split_docs = [Document(page_content=text, metadata={"source": self.file_name}) for text in split_text]
        out = chain({"input_documents": split_docs}, return_only_outputs=True)['output_text']

        # convert to json
        output = convert_json(out, item_list)

        self.article_full = {**output, **article_divided}

    def _translate_chinese(self, content:str) -> str:
        if not content: return "N/A"
        if str(detect_langs(content)[0]).split(':')[0] != self.translated_language:
            doc = Document(page_content=content, metadata={"source": self.file_name})
            prompt_template = f"You are an experienced translator who will translate the content into {self.translated_language} if the given text is not in {self.translated_language}. \
                You will translate the given text in a way that stays faithful to the original without adding much expansion and explanation. You will only return the translated text" "{text}"
        
            prompt = PromptTemplate.from_template(prompt_template)
            llm_chain = LLMChain(llm=self.llm, prompt=prompt, return_final_only=True)

            stuff_translate_chain = StuffDocumentsChain(  
                llm_chain=llm_chain, document_variable_name="text")

            return stuff_translate_chain.run([doc])
        else:
            return content

    def _get_subtopic_summary(self) -> None:
        item_list, items, item_format = get_items('individuel')

        prompt_template = f"Your primary focus should be on accurately identifying or extracting specific information.\
                            Find out or extract the {items} based on the information given in the text. \
                            Consequently, adhere to the designated format below:\
                            Subtopic:\
                            {item_format}"\
                            "{text}"

        prompt = PromptTemplate.from_template(prompt_template)

        # Define LLM chain
        llm_chain = LLMChain(llm=self.llm, prompt=prompt, return_final_only=True)

        # Define StuffDocumentsChain
        stuff_chain = StuffDocumentsChain(  
            llm_chain=llm_chain, document_variable_name="text")
        
        print('Analysing subtopics')
        result = list()
        with tqdm(total=len(self.article_full.get('Subtopics'))) as pbar:
            for subtopic in self.article_full.get('Subtopics'):
                content = f"{subtopic.get('subtopic')}: {subtopic.get('transcript').strip()}"
                doc = Document(page_content=content , metadata={"source": self.file_name})
                out = stuff_chain.run([doc])

                # convert to json
                output = convert_json(out, item_list)
                output['subtopic'] = subtopic.get('subtopic')
                output['original transcript'] = subtopic.get('transcript')
                if self.original_language != self.translated_language:
                    output['translated transcript'] = self._translate_chinese(subtopic.get('transcript'))

                if subtopic.get('timestamp'):
                    output['timestamp']= [{'start': subtopic.get('timestamp').get('start')}, {'end': subtopic.get('timestamp').get('end')}]
                result.append(output)
                pbar.update(1)
            self.article_full.update({"Subtopics":result})
        with open(f'./summary/{self.file_name}.json', 'w', encoding='utf-8') as f: json.dump(self.article_full, f, ensure_ascii=False)
        print("Analysis completed")
    
    def _translate_convert_docx(self) -> None:
        from utils import add_hyperlink, divide_audio
        import docx, datetime

        # initial a docx
        document = docx.Document()

        # translate general info and convert in docx
        items_list, _, _ = get_items('general')
        for item in items_list:
            content = self._translate_chinese(self.article_full.get(item))
            document.add_heading(item, level=1)
            document.add_paragraph(content)

        print('Translating')
        subtopics = self.article_full.get('Subtopics')
        with tqdm(total=len(subtopics)) as pbar:
            for subtopic in subtopics:
                content = self._translate_chinese(subtopic.get('subtopic'))
                insertion = document.add_heading(content, level=2)

                # add hyperlink
                if subtopic.get('timestamp') and isinstance(subtopic.get('timestamp')[0].get('start'), int) and isinstance(subtopic.get('timestamp')[1].get('end'), int):
                    start = subtopic.get('timestamp')[0].get('start')
                    end = subtopic.get('timestamp')[1].get('end')
                    subtopic_name = subtopic.get('subtopic')
                    # seperate audio by suntopics
                    absolute_path = divide_audio(self.file_name, subtopic_name, start, end)
                    add_hyperlink(insertion, f'{datetime.timedelta(seconds = int(start))}', f'file:///{absolute_path}/{subtopic.get("subtopic")}.wav')

                # translate individual item and convert in docx
                items_list, _, _ = get_items('individuel')
                for item in items_list:
                    content = self._translate_chinese(subtopic.get(item))
                    document.add_heading(item, level=3)
                    document.add_paragraph(content)

                # add chinese transcript
                if self.original_language != self.translated_language:
                    document.add_heading('translated transcript', level=3)
                    document.add_paragraph(subtopic.get("translated transcript").strip())
                document.add_heading('original transcript', level=3)
                document.add_paragraph(subtopic.get('original transcript').strip())

                document.save(f'{self.output_dir}/{self.file_name}.docx')
                pbar.update(1)

    def run(self, article_divided:dict) -> None:
        # generate global and subtopic summary
        self._get_general_summary(article_divided)
        self._get_subtopic_summary()
        
        # Translate and convert json to docx
        # with open(f'./summary/{self.file_name}.json') as f: 
        #     self.article_full = json.load(f)
        self._translate_convert_docx()
